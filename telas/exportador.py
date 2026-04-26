from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re
import base64
import urllib.request

def processar_texto_e_imagem(dado, doc, prefixo=""):
    """
    Caça imagens no HTML, limpa o texto e insere ambos no Word.
    """
    if isinstance(dado, dict):
        texto_bruto = dado.get('texto', list(dado.values())[0] if dado else "")
    else:
        texto_bruto = str(dado)
        
    # 1. Caçador de Imagens
    img_tags = re.findall(r'<img[^>]+src=["\']([^"\']+)["\']', texto_bruto)
    
    # 2. Limpa o HTML 
    texto_limpo = re.sub(r'<[^>]+>', '', texto_bruto).strip()
    
    # 3. Escreve o texto no documento 
    if texto_limpo or prefixo.strip():
        doc.add_paragraph(f"{prefixo}{texto_limpo}")
        
    # 4. Processa e desenha as imagens
    for src in img_tags:
        try:
            if src.startswith('data:image'):
                imgstr = src.split(';base64,')[1]
                img_data = base64.b64decode(imgstr)
                img_stream = BytesIO(img_data)
                doc.add_picture(img_stream, width=Inches(3.5)) 
                
            elif src.startswith('http'):
                req = urllib.request.Request(src, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req) as response:
                    img_stream = BytesIO(response.read())
                    doc.add_picture(img_stream, width=Inches(3.5))
        except Exception:
            doc.add_paragraph(" [Aviso: Uma imagem não pôde ser carregada no Word]")

def gerar_prova_word(titulo_prova, questoes):
    """
    Função (Motor) que cria o documento Word formatado.
    """
    doc = Document()
    
    # ==========================================
    # ⚙️ CONFIGURAÇÕES DE FONTE E PARÁGRAFO
    # ==========================================
    estilo_padrao = doc.styles['Normal']
    
    # Configura a Fonte (Definido como Arial, Tamanho 11)
    estilo_padrao.font.name = 'Arial'
    estilo_padrao.font.size = Pt(11)
    
    # Configura o Parágrafo (Exatamente igual à imagem enviada)
    formato_paragrafo = estilo_padrao.paragraph_format
    formato_paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Justificado
    formato_paragrafo.space_before = Pt(0) # Espaçamento Antes: 0 pt
    formato_paragrafo.space_after = Pt(0)  # Espaçamento Depois: 0 pt
    
    # ==========================================
    # 1. CABEÇALHO E QUESTÕES DA PROVA
    # ==========================================
    doc.add_heading(f'Avaliação: {titulo_prova}', 0)
    doc.add_paragraph('Nome: __________________________________________________ Turma: _______')
    doc.add_paragraph('\n') 

    for i, q in enumerate(questoes, 1):
        # Processa o enunciado
        enunciado_bruto = q.get('enunciado', '')
        processar_texto_e_imagem(enunciado_bruto, doc, prefixo=f"{i}) ")
        
        # Processa as alternativas
        opcoes = q.get('alternativas') or q.get('opcoes') or {} 
        if opcoes:
            for letra in ["A", "B", "C", "D", "E"]:
                if letra in opcoes:
                    processar_texto_e_imagem(opcoes[letra], doc, prefixo=f"    ({letra}) ")
        
        # Processa imagens em colunas avulsas (caso existam)
        url_extra = q.get('imagem_url') or q.get('imagem')
        if url_extra and str(url_extra).startswith('http'):
            try:
                req = urllib.request.Request(url_extra, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req) as response:
                    doc.add_picture(BytesIO(response.read()), width=Inches(3.5))
            except:
                pass
                
        doc.add_paragraph('\n') # Espaço entre uma questão e outra

    # ==========================================
    # 2. GABARITO (PÁGINA SEPARADA)
    # ==========================================
    doc.add_page_break() 
    doc.add_heading('Gabarito - Uso Exclusivo da Coordenação', 1)
    doc.add_paragraph('\n')

    for i, q in enumerate(questoes, 1):
        resposta_correta = q.get('resposta_correta') or q.get('gabarito') or q.get('resposta') or "Não informada"
        
        p = doc.add_paragraph()
        p.add_run(f"Questão {i}: ").bold = True
        p.add_run(f"Alternativa {resposta_correta}")

    # ==========================================
    # 3. SALVAMENTO EM MEMÓRIA
    # ==========================================
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer